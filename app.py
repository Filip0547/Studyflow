import csv
import io
import json
import os
import re
from datetime import timedelta
from functools import lru_cache
from itertools import zip_longest
from threading import Thread

import polib
from dotenv import load_dotenv
from sqlalchemy import or_
from sqlalchemy.exc import IntegrityError, SQLAlchemyError

from flask import Flask, abort, flash, jsonify, redirect, render_template, request, session, url_for
from flask_sqlalchemy import SQLAlchemy
from flask_mail import Mail, Message
from flask_babel import Babel, gettext
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.routing import BuildError
from flask_wtf import CSRFProtect
from authlib.integrations.flask_client import OAuth
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user

load_dotenv()

# Supported languages
DEFAULT_LANGUAGE = 'en'
PREFIX_LANGUAGES = ['nl', 'pl', 'es', 'fr', 'de', 'ru']
LANGUAGES = [DEFAULT_LANGUAGE] + PREFIX_LANGUAGES
LANGUAGE_NAMES = {
    'en': 'English',
    'nl': 'Nederlands',
    'pl': 'Polski',
    'es': 'Espanol',
    'fr': 'Francais',
    'de': 'Deutsch',
    'ru': 'Русский',
}

app = Flask(__name__)
# Keep SECRET_KEY stable so sessions and OAuth state survive restarts/deploys.
app.secret_key = os.getenv('SECRET_KEY', 'dev-only-insecure-secret-change-me')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024
app.config['REMEMBER_COOKIE_DURATION'] = timedelta(days=int(os.getenv('REMEMBER_COOKIE_DAYS', '30')))
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['REMEMBER_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['REMEMBER_COOKIE_SAMESITE'] = 'Lax'

# Render runs behind a proxy; trust forwarded proto/host for correct external URLs.
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)

if os.getenv('RENDER') or os.getenv('RENDER_EXTERNAL_URL'):
    app.config['SESSION_COOKIE_SECURE'] = True
    app.config['REMEMBER_COOKIE_SECURE'] = True

# Google OAuth Configuration
app.config['GOOGLE_CLIENT_ID'] = os.getenv('GOOGLE_CLIENT_ID')
app.config['GOOGLE_CLIENT_SECRET'] = os.getenv('GOOGLE_CLIENT_SECRET')

# Flask-Mail Configuration
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', True)
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER', 'noreply@studyflow.com')
app.config['MAIL_TIMEOUT'] = int(os.getenv('MAIL_TIMEOUT', 10))

# initialize extensions
db = SQLAlchemy(app)
csrf = CSRFProtect(app)
oauth = OAuth(app)
mail = Mail(app)

# Flask-Login Configuration
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

def validate_language(lang):
    """Validate language from route and return normalized value."""
    if lang is None:
        return DEFAULT_LANGUAGE
    if lang in LANGUAGES:
        return lang
    abort(404)


def get_current_language():
    """Language is derived from route path, never from session."""
    view_args = request.view_args or {}
    lang = view_args.get('lang')
    if lang in PREFIX_LANGUAGES:
        return lang
    if lang == DEFAULT_LANGUAGE:
        return DEFAULT_LANGUAGE
    return DEFAULT_LANGUAGE


def get_locale():
    """Locale selector used by Flask-Babel, aligned with route language."""
    return get_current_language()

babel = Babel(app, locale_selector=get_locale)


@lru_cache(maxsize=64)
def load_po_catalog(lang):
    """Load a msgid->msgstr dict from translations/<lang>/LC_MESSAGES/messages.po."""
    po_path = os.path.join(app.root_path, 'translations', lang, 'LC_MESSAGES', 'messages.po')
    if not os.path.exists(po_path):
        return {}

    try:
        catalog = {}
        for entry in polib.pofile(po_path):
            if entry.obsolete or not entry.msgid:
                continue
            if entry.msgstr:
                catalog[entry.msgid] = entry.msgstr
        return catalog
    except Exception:
        return {}


def translate_text(message, *args, **kwargs):
    """Translate with Babel first, then direct .po fallback for robustness."""
    lang = get_current_language()

    translated = message
    try:
        babel_value = gettext(message)
        if babel_value and babel_value != message:
            translated = babel_value
        else:
            translated = load_po_catalog(lang).get(message, message)
    except Exception:
        translated = load_po_catalog(lang).get(message, message)

    try:
        if kwargs:
            translated = translated % kwargs
        elif args:
            translated = translated % args
    except Exception:
        pass

    return translated


def localized_url(endpoint, lang=None, **values):
    """Build URLs that preserve language prefix (except English canonical routes)."""
    target_lang = validate_language(lang) if lang else get_current_language()
    cleaned_values = {k: v for k, v in values.items() if v is not None}

    if target_lang == DEFAULT_LANGUAGE:
        cleaned_values.pop('lang', None)
    else:
        cleaned_values['lang'] = target_lang

    return url_for(endpoint, **cleaned_values)


def switch_language_url(lang):
    """Build the current page URL in another language."""
    target_lang = validate_language(lang)
    endpoint = request.endpoint

    if not endpoint or endpoint == 'static':
        return localized_url('index', lang=target_lang)

    args = dict(request.view_args or {})
    args.pop('lang', None)

    try:
        base_url = localized_url(endpoint, lang=target_lang, **args)
    except BuildError:
        base_url = localized_url('index', lang=target_lang)

    if request.query_string:
        return f"{base_url}?{request.query_string.decode('utf-8')}"
    return base_url


def active_page():
    """Return endpoint name for active nav state."""
    return request.endpoint or ''


def parse_initial_words(raw_words):
    """Parse textarea lines in `word|description|example|disadvantage` format."""
    parsed_words = []
    if not raw_words:
        return parsed_words

    for line in raw_words.splitlines():
        cleaned = line.strip()
        if not cleaned:
            continue

        parts = [part.strip() for part in cleaned.split('|')]
        if not parts[0]:
            continue

        while len(parts) < 4:
            parts.append('')

        parsed_words.append(
            {
                'word': parts[0],
                'description': parts[1],
                'example': parts[2],
                'disadvantage': parts[3],
            }
        )

    return parsed_words


def normalize_word_entry(word='', description='', example='', disadvantage=''):
    """Return a normalized word payload for forms, imports, and study views."""
    return {
        'word': (word or '').strip(),
        'description': (description or '').strip(),
        'example': (example or '').strip(),
        'disadvantage': (disadvantage or '').strip(),
    }


def normalize_word_source(source):
    """Normalize either a dict payload or a Word model instance."""
    if isinstance(source, dict):
        return normalize_word_entry(
            source.get('word', ''),
            source.get('description', ''),
            source.get('example', ''),
            source.get('disadvantage', ''),
        )

    return normalize_word_entry(
        getattr(source, 'word', ''),
        getattr(source, 'description', ''),
        getattr(source, 'example', ''),
        getattr(source, 'disadvantage', ''),
    )


def build_editor_rows(entries, minimum_rows=3):
    """Prepare rows for the bulk editor and guarantee some empty starter rows."""
    rows = [normalize_word_source(entry) for entry in entries]
    while len(rows) < minimum_rows:
        rows.append(normalize_word_entry())
    return rows


def extract_word_rows_from_form(form):
    """Read bulk editor form rows and ignore completely empty or wordless rows."""
    columns = [
        form.getlist('word[]'),
        form.getlist('description[]'),
        form.getlist('example[]'),
        form.getlist('disadvantage[]'),
    ]
    rows = []

    for word, description, example, disadvantage in zip_longest(*columns, fillvalue=''):
        row = normalize_word_entry(word, description, example, disadvantage)
        if not any(row.values()):
            continue
        if not row['word']:
            continue
        rows.append(row)

    return rows


def build_flashcard_rows(words):
    """Return saved rows that contain at least one study field."""
    rows = []
    for word in words:
        row = normalize_word_source(word)
        if not row['word']:
            continue
        if not any([row['description'], row['example'], row['disadvantage']]):
            continue
        rows.append(row)
    return rows


def build_quiz_rows(words):
    """Return rows that can be used in description-based quizzes."""
    return [row for row in build_flashcard_rows(words) if row['description']]


def build_learn_rows(words):
    """Return rows suitable for Learn Mode (term + definition required)."""
    rows = []
    for source in words:
        row = normalize_word_source(source)
        if not row['word'] or not row['description']:
            continue
        rows.append(
            {
                'id': getattr(source, 'id', len(rows) + 1),
                'word': row['word'],
                'description': row['description'],
            }
        )
    return rows


def decode_text_bytes(file_bytes):
    """Decode uploaded text using a few common encodings."""
    for encoding in ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1'):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode('utf-8', errors='ignore')


def detect_separator(lines, preferred='auto'):
    """Choose the best separator for pasted or extracted text."""
    explicit_map = {
        'tab': '\t',
        'equals': '=',
        'pipe': '|',
        'semicolon': ';',
        'comma': ',',
    }
    if preferred in explicit_map:
        return explicit_map[preferred]

    candidates = ['\t', '|', '=', ';', ',']
    separator_scores = {}
    for separator in candidates:
        separator_scores[separator] = sum(1 for line in lines if separator in line)

    best_separator = max(separator_scores, key=separator_scores.get, default='')
    if separator_scores.get(best_separator, 0) > 0:
        return best_separator

    if any(re.search(r'\S\s{2,}\S', line) for line in lines):
        return 'double-space'

    return ''


def parse_import_text(raw_text, separator='auto'):
    """Parse pasted text or extracted document text into editor rows."""
    if not raw_text:
        return []

    lines = [line.strip() for line in raw_text.replace('\r\n', '\n').split('\n') if line.strip()]
    if not lines:
        return []

    active_separator = detect_separator(lines, preferred=separator)
    rows = []

    def parse_single_line_payload(line_text):
        """Parse common one-line formats like `word - description` or `word: description`."""
        cleaned = re.sub(r'^\s*[-*•]\s*', '', line_text.strip())
        if not cleaned:
            return ('', '')

        match = re.match(r'^(?P<word>.+?)\s*(?:->|=>|:|=|\s[-–—]\s)\s*(?P<description>.+)$', cleaned)
        if match:
            return (match.group('word').strip(), match.group('description').strip())

        return (cleaned, '')

    if active_separator:
        for line in lines:
            if active_separator == 'double-space':
                parts = [part.strip() for part in re.split(r'\s{2,}', line) if part.strip()]
            else:
                if active_separator not in line:
                    continue
                parts = [part.strip() for part in line.split(active_separator)]

            if not parts:
                continue

            if not parts[0] and len(parts) > 1 and parts[1]:
                parts = parts[1:]

            if len(parts) == 1:
                parsed_word, parsed_description = parse_single_line_payload(parts[0])
                parts = [parsed_word, parsed_description, '', '']

            while len(parts) < 4:
                parts.append('')
            if len(parts) > 4:
                parts = parts[:3] + [' '.join(part for part in parts[3:] if part).strip()]

            row = normalize_word_entry(parts[0], parts[1], parts[2], parts[3])
            if row['word']:
                rows.append(row)

    if rows:
        return rows

    fallback_rows = []
    index = 0
    while index < len(lines):
        parsed_word, parsed_description = parse_single_line_payload(lines[index])

        if parsed_description:
            row = normalize_word_entry(parsed_word, parsed_description)
            if row['word']:
                fallback_rows.append(row)
            index += 1
            continue

        description = lines[index + 1] if index + 1 < len(lines) else ''
        row = normalize_word_entry(parsed_word, description)
        if row['word']:
            fallback_rows.append(row)
        index += 2

    return fallback_rows


def extract_text_from_txt(file_bytes):
    """Extract text from plain-text uploads."""
    return decode_text_bytes(file_bytes)


def extract_text_from_csv(file_bytes, delimiter=','):
    """Extract rows from CSV-like files as tab-separated lines."""
    decoded = decode_text_bytes(file_bytes)
    reader = csv.reader(io.StringIO(decoded), delimiter=delimiter)
    output_lines = []
    for row in reader:
        cleaned = [str(cell).strip() for cell in row]
        if any(cleaned):
            output_lines.append('\t'.join(cleaned))
    return '\n'.join(output_lines)


def extract_text_from_docx(file_bytes):
    """Extract paragraphs and table rows from a .docx document."""
    from docx import Document

    document = Document(io.BytesIO(file_bytes))
    chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append('\t'.join(cells))

    return '\n'.join(chunks)


def extract_text_from_pdf(file_bytes):
    """Extract text from PDF pages."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    chunks = []
    for page in reader.pages:
        page_text = (page.extract_text() or '').strip()
        if page_text:
            chunks.append(page_text)
    return '\n'.join(chunks)


def extract_text_from_xlsx(file_bytes):
    """Extract rows from spreadsheets as tab-separated lines."""
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(file_bytes), data_only=True)
    output_lines = []
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows(values_only=True):
            cleaned = [str(cell).strip() if cell is not None else '' for cell in row]
            if any(cleaned):
                output_lines.append('\t'.join(cleaned))
    return '\n'.join(output_lines)


@lru_cache(maxsize=1)
def get_ocr_engine():
    """Create the OCR engine lazily so image import stays optional until needed."""
    from rapidocr_onnxruntime import RapidOCR

    return RapidOCR()


def extract_text_from_image(file_bytes):
    """Extract OCR text from image uploads."""
    import numpy as np
    from PIL import Image

    image = Image.open(io.BytesIO(file_bytes)).convert('RGB')
    result, _ = get_ocr_engine()(np.array(image))
    if not result:
        return ''
    return '\n'.join(item[1].strip() for item in result if len(item) > 1 and item[1].strip())


def extract_text_from_upload(uploaded_file):
    """Extract plain text from supported uploads."""
    filename = (uploaded_file.filename or '').strip()
    if not filename:
        return ''

    extension = os.path.splitext(filename)[1].lower()
    file_bytes = uploaded_file.read()
    if not file_bytes:
        return ''

    if extension == '.txt':
        return extract_text_from_txt(file_bytes)
    if extension == '.csv':
        return extract_text_from_csv(file_bytes, delimiter=',')
    if extension == '.tsv':
        return extract_text_from_csv(file_bytes, delimiter='\t')
    if extension == '.docx':
        return extract_text_from_docx(file_bytes)
    if extension == '.pdf':
        return extract_text_from_pdf(file_bytes)
    if extension == '.xlsx':
        return extract_text_from_xlsx(file_bytes)
    if extension in {'.png', '.jpg', '.jpeg', '.webp', '.bmp'}:
        return extract_text_from_image(file_bytes)

    raise ValueError('Unsupported file type.')


def asset_url(filename):
    """Static URL with per-file cache busting."""
    full_path = os.path.join(app.static_folder, filename)
    version = int(os.path.getmtime(full_path)) if os.path.exists(full_path) else 0
    return url_for('static', filename=filename, v=version)

# Register Google OAuth
google = None
if app.config['GOOGLE_CLIENT_ID'] and app.config['GOOGLE_CLIENT_SECRET']:
    google = oauth.register(
        name='google',
        client_id=app.config['GOOGLE_CLIENT_ID'],
        client_secret=app.config['GOOGLE_CLIENT_SECRET'],
        server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
        client_kwargs={'scope': 'openid email profile'},
    )


@login_manager.user_loader
def load_user(user_id):
    """Load user by ID for Flask-Login."""
    return User.query.get(int(user_id))


# ─── MODELS ───────────────────────────────────────────────

class User(UserMixin, db.Model):
    id       = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=True)
    email    = db.Column(db.String(100), unique=True, nullable=True)
    password = db.Column(db.String(200), nullable=True)
    google_id = db.Column(db.String(100), unique=True, nullable=True)
    lists    = db.relationship('WordList', backref='owner', lazy=True)


class WordList(db.Model):
    id      = db.Column(db.Integer, primary_key=True)
    name    = db.Column(db.String(100), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    words   = db.relationship('Word', backref='word_list', lazy=True, cascade="all, delete-orphan")


class Word(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    word        = db.Column(db.String(100), nullable=False)
    description = db.Column(db.String(500))
    example     = db.Column(db.String(500))
    disadvantage= db.Column(db.String(500))
    list_id     = db.Column(db.Integer, db.ForeignKey('word_list.id'), nullable=False)


class LearnProgress(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    list_id = db.Column(db.Integer, db.ForeignKey('word_list.id'), nullable=False)
    state_json = db.Column(db.Text, nullable=False, default='{}')
    updated_at = db.Column(db.DateTime, nullable=False, server_default=db.func.now(), onupdate=db.func.now())

    __table_args__ = (
        db.UniqueConstraint('user_id', 'list_id', name='uq_learn_progress_user_list'),
    )


with app.app_context():
    db.create_all()


# ─── HELPERS ──────────────────────────────────────────────

def get_current_user():
    """Returns the logged-in User object, or None."""
    return current_user if current_user.is_authenticated else None


def send_welcome_email(user_email, username):
    """Send a welcome email to the new user."""
    try:
        msg = Message(
            subject='Welcome to StudyFlow',
            recipients=[user_email],
            html=f"""
            <html>
                <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
                    <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
                        <div style="background-color: #4f46e5; color: white; padding: 20px; border-radius: 8px; text-align: center; margin-bottom: 20px;">
                            <h1 style="margin: 0; font-size: 28px;">📚 Welcome to StudyFlow</h1>
                        </div>
                        
                        <p>Hi <strong>{username}</strong>,</p>
                        
                        <p>Thank you for joining StudyFlow! We're excited to help you learn faster and smarter.</p>
                        
                        <p>With StudyFlow, you can:</p>
                        <ul>
                            <li>Create custom study lists</li>
                            <li>Organize words and their meanings</li>
                            <li>Track your learning progress</li>
                            <li>Study efficiently with our interactive tools</li>
                        </ul>
                        
                        <p style="margin-top: 30px;">
                            <a href="https://studyflow.onrender.com/dashboard" style="display: inline-block; background-color: #4f46e5; color: white; padding: 12px 30px; text-decoration: none; border-radius: 6px; font-weight: bold;">Start Studying Now</a>
                        </p>
                        
                        <p style="margin-top: 30px; color: #666; font-size: 14px;">
                            If you have any questions, feel free to reply to this email.
                        </p>
                        
                        <hr style="border: none; border-top: 1px solid #e0e0e0; margin: 30px 0;">
                        
                        <p style="color: #999; font-size: 12px; text-align: center;">
                            StudyFlow &nbsp; | &nbsp; Learn Smarter, Not Harder
                        </p>
                    </div>
                </body>
            </html>
            """
        )
        mail.send(msg)
        return True
    except Exception as e:
        print(f"Error sending email to {user_email}: {str(e)}")
        return False


def send_welcome_email_async(user_email, username):
    """Send welcome email in the background so registration returns quickly."""

    def _send():
        with app.app_context():
            send_welcome_email(user_email, username)

    Thread(target=_send, daemon=True).start()


# ─── CONTEXT PROCESSOR ────────────────────────────────────
# Makes `current_user` and `get_locale` available in ALL templates automatically

@app.context_processor
def inject_globals():
    return dict(
        current_user=get_current_user(),
        get_locale=get_locale,
        get_current_language=get_current_language,
        localized_url=localized_url,
        switch_language_url=switch_language_url,
        active_page=active_page,
        language_names=LANGUAGE_NAMES,
        languages=LANGUAGES,
        default_language=DEFAULT_LANGUAGE,
        google_login_enabled=google is not None,
        asset_url=asset_url,
        _=translate_text,
    )


@login_manager.unauthorized_handler
def unauthorized():
    flash(translate_text('Please log in to continue.'), 'error')
    return redirect(localized_url('login', lang=get_current_language()))


@app.errorhandler(SQLAlchemyError)
def handle_db_error(error):
    db.session.rollback()
    flash(translate_text('A database error occurred. Please try again.'), 'error')
    return redirect(localized_url('index', lang=get_current_language()))


@app.errorhandler(500)
def handle_internal_error(error):
    flash(translate_text('Something went wrong. Please try again.'), 'error')
    return redirect(localized_url('index', lang=get_current_language()))


# ─── ROUTES ───────────────────────────────────────────────

@app.route("/", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>")
def index(lang):
    lang = validate_language(lang)
    # Auto redirect to dashboard if already logged in
    if current_user.is_authenticated:
        return redirect(localized_url('dashboard', lang=lang))
    return render_template("index.html")


@app.route("/register", defaults={'lang': DEFAULT_LANGUAGE}, methods=["GET", "POST"])
@app.route("/<lang>/register", methods=["GET", "POST"])
def register(lang):
    lang = validate_language(lang)
    if current_user.is_authenticated:
        return redirect(localized_url('dashboard', lang=lang))

    if request.method == "POST":
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        accepted_privacy = request.form.get('accept_privacy') == 'on'

        # Validation errors
        if not username or not email or not password:
            flash(translate_text('All fields are required.'), 'error')
        elif not accepted_privacy:
            flash(translate_text('You must accept the privacy policy to create an account.'), 'error')
        elif len(password) < 4:
            flash(translate_text('Password must be at least 4 characters long.'), 'error')
        elif not ('@' in email and '.' in email):
            flash(translate_text('Please enter a valid email address.'), 'error')
        elif User.query.filter_by(username=username).first():
            flash(translate_text('Username is already taken.'), 'error')
        elif User.query.filter_by(email=email).first():
            flash(translate_text('Email is already registered.'), 'error')
        else:
            # Create new user with email
            hashed = generate_password_hash(password)
            new_user = User(
                username=username,
                email=email,
                password=hashed
            )
            try:
                db.session.add(new_user)
                db.session.commit()
            except IntegrityError:
                db.session.rollback()
                flash(translate_text('Username or email is already in use.'), 'error')
                return render_template("register.html")
            except SQLAlchemyError:
                db.session.rollback()
                flash(translate_text('Could not create your account right now. Please try again.'), 'error')
                return render_template("register.html")

            send_welcome_email_async(email, username)
            login_user(new_user, remember=True)
            flash(translate_text('Account created successfully! You are now signed in.'), 'success')
            return redirect(localized_url('dashboard', lang=lang))

    return render_template("register.html")


@app.route("/login", defaults={'lang': DEFAULT_LANGUAGE}, methods=["GET", "POST"])
@app.route("/<lang>/login", methods=["GET", "POST"])
def login(lang):
    lang = validate_language(lang)
    if current_user.is_authenticated:
        return redirect(localized_url('dashboard', lang=lang))

    if request.method == "POST":
        identifier = request.form.get('username', '').strip()
        password = request.form.get('password', '')

        user = User.query.filter(
            or_(User.username == identifier, User.email == identifier)
        ).first()

        if user and user.password and check_password_hash(user.password, password):
            login_user(user, remember=True)
            return redirect(localized_url('dashboard', lang=lang))
        elif user and not user.password:
            flash(translate_text('This account uses Google sign-in. Please continue with Google.'), 'error')
        else:
            flash(translate_text('Invalid username/email or password.'), 'error')

    return render_template("login.html")


@app.route("/login/google", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/login/google")
def login_google(lang):
    lang = validate_language(lang)
    """Redirect user to the Google OAuth login page."""
    if google is None:
        flash(translate_text('Google login is not configured yet. Please use username/password login.'), 'error')
        return redirect(localized_url('login', lang=lang))

    session['oauth_lang'] = lang
    redirect_uri = os.getenv('GOOGLE_REDIRECT_URI') or url_for('auth_google_callback_entry', _external=True)
    return google.authorize_redirect(redirect_uri)


@app.route("/auth/google/callback")
def auth_google_callback_entry():
    lang = validate_language(session.get('oauth_lang', DEFAULT_LANGUAGE))
    return auth_google_callback(lang)


@app.route("/login/google/authorized", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/login/google/authorized")
def auth_google_callback(lang):
    lang = validate_language(lang)
    target_lang = validate_language(session.pop('oauth_lang', lang))
    """Handle the Google OAuth callback."""
    if google is None:
        flash(translate_text('Google login is not configured yet. Please use username/password login.'), 'error')
        return redirect(localized_url('login', lang=target_lang))

    try:
        # Retrieve the OAuth access token from Google
        token = google.authorize_access_token()
    except Exception:
        flash(translate_text('Failed to authorize with Google.'), 'error')
        return redirect(localized_url('login', lang=target_lang))
    
    try:
        # Fetch the user's Google profile information
        resp = google.get('userinfo')
        user_info = resp.json()
    except Exception:
        flash(translate_text('Could not retrieve user information from Google.'), 'error')
        return redirect(localized_url('login', lang=target_lang))
    
    # Extract email and name from user info
    email = user_info.get("email")
    name = user_info.get("name")
    google_id = user_info.get("sub")  # Google's unique user ID
    
    if not email or not google_id:
        flash(translate_text('Google account missing required information.'), 'error')
        return redirect(localized_url('login', lang=target_lang))
    
    # Check if user already exists by Google ID
    user = User.query.filter_by(google_id=google_id).first()
    
    if not user:
        # Check if email is already registered with username/password
        user = User.query.filter_by(email=email).first()
        
        if not user:
            safe_username = email.split('@')[0]
            base_username = safe_username or 'google_user'
            candidate = base_username
            suffix = 1
            while User.query.filter_by(username=candidate).first():
                suffix += 1
                candidate = f"{base_username}{suffix}"

            # Create a new user account
            user = User(
                email=email,
                google_id=google_id,
                username=candidate,
                password=None   # No password for Google OAuth users
            )
            try:
                db.session.add(user)
                db.session.commit()
            except SQLAlchemyError:
                db.session.rollback()
                flash(translate_text('Could not complete Google sign-in. Please try again.'), 'error')
                return redirect(localized_url('login', lang=target_lang))
            flash(translate_text('Account created successfully! You are now signed in with Google.'), 'success')
        else:
            user.google_id = user.google_id or google_id
            try:
                db.session.commit()
            except SQLAlchemyError:
                db.session.rollback()
                flash(translate_text('Could not complete Google sign-in. Please try again.'), 'error')
                return redirect(localized_url('login', lang=target_lang))
    
    # Log the user in using Flask-Login
    login_user(user, remember=True)
    return redirect(localized_url('dashboard', lang=target_lang))


@app.route("/logout", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/logout")
def logout(lang):
    lang = validate_language(lang)
    """Log out the current user."""
    logout_user()
    return redirect(localized_url('index', lang=lang))


@app.route("/contact", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/contact")
def contact(lang):
    validate_language(lang)
    return render_template("contact.html")


@app.route("/privacy-policy", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/privacy-policy")
def privacy_policy(lang):
    validate_language(lang)
    return render_template("privacy_policy.html")


@app.route("/dashboard", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/dashboard")
@login_required
def dashboard(lang):
    validate_language(lang)
    user = get_current_user()
    lists = WordList.query.filter_by(user_id=user.id).order_by(WordList.id.desc()).all()
    return render_template("dashboard.html", lists=lists)


@app.route("/create_list", defaults={'lang': DEFAULT_LANGUAGE}, methods=["GET", "POST"])
@app.route("/<lang>/create_list", methods=["GET", "POST"])
@login_required
def create_list(lang):
    lang = validate_language(lang)
    if request.method == "POST":
        name = request.form.get('name', '').strip()
        if not name:
            flash(translate_text('List name is required.'), 'error')
        else:
            new_list = WordList(name=name, user_id=current_user.id)
            db.session.add(new_list)
            db.session.commit()
            flash(translate_text('List created successfully.'), 'success')
            return redirect(localized_url('edit_list', lang=lang, list_id=new_list.id))

    return render_template("create_list.html")


@app.route("/list/<int:list_id>/edit", defaults={'lang': DEFAULT_LANGUAGE}, methods=["GET", "POST"])
@app.route("/<lang>/list/<int:list_id>/edit", methods=["GET", "POST"])
@login_required
def edit_list(list_id, lang):
    lang = validate_language(lang)
    user = get_current_user()
    word_list = WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()
    saved_words = Word.query.filter_by(list_id=list_id).order_by(Word.id.asc()).all()

    if request.method == "POST":
        name = request.form.get('name', '').strip()
        editor_rows = extract_word_rows_from_form(request.form)

        if not name:
            flash(translate_text('List name is required.'), 'error')
            return render_template(
                "edit_list.html",
                word_list=word_list,
                editor_rows=build_editor_rows(editor_rows),
                saved_count=len(saved_words),
            )

        word_list.name = name
        Word.query.filter_by(list_id=list_id).delete(synchronize_session=False)
        for row in editor_rows:
            db.session.add(
                Word(
                    word=row['word'],
                    description=row['description'],
                    example=row['example'],
                    disadvantage=row['disadvantage'],
                    list_id=list_id,
                )
            )

        db.session.commit()
        flash(translate_text('List saved successfully.'), 'success')
        return redirect(localized_url('edit_list', lang=lang, list_id=list_id))

    return render_template(
        "edit_list.html",
        word_list=word_list,
        editor_rows=build_editor_rows(saved_words),
        saved_count=len(saved_words),
    )


@app.route("/list/<int:list_id>/import_words", defaults={'lang': DEFAULT_LANGUAGE}, methods=["POST"])
@app.route("/<lang>/list/<int:list_id>/import_words", methods=["POST"])
@login_required
def import_words(list_id, lang):
    lang = validate_language(lang)
    user = get_current_user()
    WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()

    separator = request.form.get('separator', 'auto').strip() or 'auto'
    paste_text = request.form.get('paste_text', '').strip()
    uploaded_file = request.files.get('source_file')
    imported_rows = []

    if paste_text:
        imported_rows.extend(parse_import_text(paste_text, separator=separator))

    if uploaded_file and uploaded_file.filename:
        try:
            extracted_text = extract_text_from_upload(uploaded_file)
        except ValueError:
            return jsonify({'error': translate_text('Unsupported file type. Use txt, csv, tsv, docx, xlsx, pdf, or an image.')}), 400
        except Exception:
            return jsonify({'error': translate_text('The file could not be read. Try a clearer image or a simpler document.')}), 400

        imported_rows.extend(parse_import_text(extracted_text, separator=separator))

    if not imported_rows:
        return jsonify({'error': translate_text('No words were found in the pasted text or uploaded file.')}), 400

    return jsonify(
        {
            'message': translate_text('%(count)s rows imported into the editor.', count=len(imported_rows)),
            'rows': imported_rows,
        }
    )


@app.route("/list/<int:list_id>/flashcards", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/list/<int:list_id>/flashcards")
@login_required
def flashcards(list_id, lang):
    validate_language(lang)
    user = get_current_user()
    word_list = WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()
    words = build_flashcard_rows(Word.query.filter_by(list_id=list_id).all())
    return render_template("flashcards.html", word_list=word_list, words=words, words_data=words)


@app.route("/list/<int:list_id>/quiz", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/list/<int:list_id>/quiz")
@login_required
def quiz(list_id, lang):
    validate_language(lang)
    user = get_current_user()
    word_list = WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()
    study_words = build_quiz_rows(Word.query.filter_by(list_id=list_id).all())
    return render_template("quiz.html", word_list=word_list, words=study_words, words_data=study_words)


@app.route("/list/<int:list_id>/multiple-choice", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/list/<int:list_id>/multiple-choice")
@login_required
def multiple_choice(list_id, lang):
    validate_language(lang)
    user = get_current_user()
    word_list = WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()
    study_words = build_quiz_rows(Word.query.filter_by(list_id=list_id).all())
    return render_template("multiple_choice.html", word_list=word_list, words=study_words, words_data=study_words)


@app.route("/list/<int:list_id>/learn", defaults={'lang': DEFAULT_LANGUAGE})
@app.route("/<lang>/list/<int:list_id>/learn")
@login_required
def learn_mode(list_id, lang):
    validate_language(lang)
    user = get_current_user()
    word_list = WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()
    study_words = build_learn_rows(Word.query.filter_by(list_id=list_id).all())
    progress_row = LearnProgress.query.filter_by(user_id=user.id, list_id=list_id).first()

    initial_state = None
    if progress_row and progress_row.state_json:
        try:
            loaded_state = json.loads(progress_row.state_json)
            if isinstance(loaded_state, dict):
                initial_state = loaded_state
        except (TypeError, ValueError):
            initial_state = None

    return render_template(
        "learn_mode.html",
        word_list=word_list,
        words=study_words,
        words_data=study_words,
        initial_state=initial_state,
    )


@app.route("/list/<int:list_id>/learn-progress", defaults={'lang': DEFAULT_LANGUAGE}, methods=["POST"])
@app.route("/<lang>/list/<int:list_id>/learn-progress", methods=["POST"])
@login_required
def save_learn_progress(list_id, lang):
    validate_language(lang)
    user = get_current_user()
    WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()

    payload = request.get_json(silent=True) or {}
    state = payload.get('state')
    if not isinstance(state, dict):
        return jsonify({'error': translate_text('Invalid progress payload.')}), 400

    allowed_phases = {'reading', 'multiple_choice', 'typing_hint', 'typing_plain', 'completed'}
    if state.get('phase') not in allowed_phases:
        return jsonify({'error': translate_text('Invalid progress payload.')}), 400

    progress_row = LearnProgress.query.filter_by(user_id=user.id, list_id=list_id).first()
    if not progress_row:
        progress_row = LearnProgress(user_id=user.id, list_id=list_id)
        db.session.add(progress_row)

    progress_row.state_json = json.dumps(state)
    db.session.commit()
    return jsonify({'ok': True})


@app.route("/list/<int:list_id>/learn-progress/reset", defaults={'lang': DEFAULT_LANGUAGE}, methods=["POST"])
@app.route("/<lang>/list/<int:list_id>/learn-progress/reset", methods=["POST"])
@login_required
def reset_learn_progress(list_id, lang):
    validate_language(lang)
    user = get_current_user()
    WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()

    progress_row = LearnProgress.query.filter_by(user_id=user.id, list_id=list_id).first()
    if progress_row:
        db.session.delete(progress_row)
        db.session.commit()

    return jsonify({'ok': True})


@app.route("/list/<int:list_id>/delete_word/<int:word_id>", defaults={'lang': DEFAULT_LANGUAGE}, methods=["POST"])
@app.route("/<lang>/list/<int:list_id>/delete_word/<int:word_id>", methods=["POST"])
@login_required
def delete_word(list_id, word_id, lang):
    lang = validate_language(lang)
    user = get_current_user()
    # Make sure the list belongs to this user
    WordList.query.filter_by(id=list_id, user_id=user.id).first_or_404()
    word = Word.query.filter_by(id=word_id, list_id=list_id).first_or_404()
    db.session.delete(word)
    db.session.commit()
    return redirect(localized_url('edit_list', lang=lang, list_id=list_id))


@app.route("/set_language/<lang>")
def set_language(lang):
    """Legacy language switch endpoint kept for backward compatibility."""
    if lang not in LANGUAGES:
        lang = DEFAULT_LANGUAGE
    return redirect(localized_url('index', lang=lang))


if __name__ == "__main__":
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', '').lower() in {'1', 'true', 'yes', 'on'}
    app.run(host='0.0.0.0', port=port, debug=debug)